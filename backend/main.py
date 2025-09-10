from io import BytesIO
from fastapi import FastAPI, File, UploadFile, Form
from fastapi.responses import JSONResponse, StreamingResponse, FileResponse
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from fastapi.requests import Request
import uvicorn
import json
import logging
import os
import shutil
import secrets
import fitz
from tempfile import NamedTemporaryFile
from typing import Optional
from dotenv import load_dotenv
from starlette.middleware.base import BaseHTTPMiddleware


import google.generativeai as genai

import docx
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from PyPDF2 import PdfReader
templates = Jinja2Templates(directory="public")

# Load environment variables
load_dotenv()
genai.configure(api_key=os.getenv("GEMINI_API_KEY"))

app = FastAPI(title="AI Resume & Career Guidance Bot")
logging.basicConfig(level=logging.INFO)  # Set level to DEBUG, INFO, etc.
logger = logging.getLogger(__name__)
# Enable CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)
class CSPNonceMiddleware(BaseHTTPMiddleware):
    async def dispatch(self, request, call_next):
        nonce = secrets.token_urlsafe(16)
        request.state.nonce = nonce  # ✅ store in request state

        response = await call_next(request)

        csp_value = (
            f"default-src 'self'; "
            f"connect-src 'self' https://ai-resume-generator-rw01.onrender.com; "
            f"script-src 'self' 'nonce-{nonce}'; "
            f"style-src 'self' 'nonce-{nonce}'; "
            f"img-src 'self' data:;"
        )
        response.headers["Content-Security-Policy"] = csp_value
        response.headers["X-CSP-Nonce"] = nonce

        return response

app.add_middleware(CSPNonceMiddleware)
app.mount("/static", StaticFiles(directory="public"), name="static")

@app.get("/")
async def serve_index(request: Request):
    # Get nonce from middleware
    nonce = getattr(request.state, "nonce", "")
    return templates.TemplateResponse("index.html", {"request": request, "nonce": nonce})


# Utility functions
def extract_text_from_pdf(file_path: str) -> str:
    text = []
    with open(file_path, "rb") as fh:
        reader = PdfReader(fh)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text.append(page_text)
    return "\n".join(text)

def extract_text_from_docx(file_path: str) -> str:
    doc = docx.Document(file_path)
    return "\n".join([p.text for p in doc.paragraphs])

async def gemini_call(prompt: str, model: str = "gemini-2.5-pro") -> str:
    try:
        response = genai.GenerativeModel(model).generate_content(prompt)
        return response.text if response and response.text else "No response generated."
    except Exception as e:
        return f"Gemini API error: {str(e)}"

def parse_resume_with_ai(text: str):
    prompt = f"""
    Extract a JSON from the following resume text.
    Format strictly as JSON with fields:
    {{
      "summary": "",
      "skills": ["skill1", "skill2"],
      "experience": [{{"title":"", "company":"", "duration":"", "description":""}}],
      "education": [{{"degree":"", "institution":"", "year":""}}]
    }}
    Resume Text:
    {text}
    """
    try:
        ai_text = genai.GenerativeModel("gemini-2.5-pro").generate_content(prompt).text
        data = json.loads(ai_text)
    except:
        data = {"summary": text[:200], "skills": [], "experience": [], "education": []}
    return data

# --- Endpoints ---

@app.get("/hello")
def hello():
    logger.info("Hello endpoint was called")
    try:
        # your logic
        logger.debug("Additional debug info")
    except Exception as e:
        logger.error(f"Error occurred: {e}", exc_info=True)
    return {"message": "Hello World"}


@app.post("/analyze")
async def analyze_resume(
    text: Optional[str] = Form(None),
    file: Optional[UploadFile] = File(None)
):
    extracted_text = ""

    # 1️⃣ Extract text from uploaded file
    if file:
        suffix = os.path.splitext(file.filename)[1].lower()
        with NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
            tmp_name = tmp.name
            shutil.copyfileobj(file.file, tmp)
        try:
            if suffix == ".pdf":
                extracted_text = extract_text_from_pdf(tmp_name)
            elif suffix in [".docx", ".doc"]:
                extracted_text = extract_text_from_docx(tmp_name)
            else:
                with open(tmp_name, "r", encoding="utf-8", errors="ignore") as fh:
                    extracted_text = fh.read()
        finally:
            os.remove(tmp_name)
    elif text:
        extracted_text = text
    else:
        return JSONResponse({"error": "No text or file provided."}, status_code=400)

    # 2️⃣ Parse with AI and ensure JSON fields
    try:
        prompt = f"""
        You are an AI that extracts structured resume data.

        Extract JSON with fields:
        {{
          "summary": "A 2-3 sentence professional summary",
          "skills": ["list of technical or soft skills"],
          "experience": [{"title":"", "company":"", "duration":"", "description":""}],
          "education": [{"degree":"", "institution":"", "year":""}]
        }}

        Resume Text:
        {extracted_text}

        ONLY return valid JSON.
        """

        ai_text = genai.GenerativeModel("gemini-2.5-pro").generate_content(prompt).text
        data = json.loads(ai_text)

        # 3️⃣ Ensure summary and skills are not empty
        if not data.get("summary"):
            data["summary"] = "\n".join(extracted_text.split("\n")[:3])

        if not data.get("skills") or not isinstance(data["skills"], list):
            # Simple fallback: extract keywords from resume text
            text_lower = extracted_text.lower()
            common_skills = ["python", "java", "javascript", "c++", "sql", "fastapi", "react", "html", "css"]
            data["skills"] = [s.capitalize() for s in common_skills if s in text_lower]

    except Exception as e:
        # Fallback if AI fails completely
        print("AI parsing failed:", e)
        data = {
            "summary": "\n".join(extracted_text.split("\n")[:3]),
            "skills": [],
            "experience": [],
            "education": []
        }

    # 4️⃣ Return JSON
    return {
        "extracted_text_snippet": extracted_text[:300],
        **data  # summary, skills, experience, education
    }


@app.post("/generate")
async def generate_resume(
    name: str = Form(...),
    email: str = Form(...),
    phone: str = Form(...),
    location: str = Form(...),
    summary: str = Form(""),
    skills: str = Form(""),
    education_json: str = Form("[]"),
    experience_json: str = Form("[]"),
):
    try:
        education = json.loads(education_json)
        experience = json.loads(experience_json)

        doc = Document()

        # Header
        header = doc.add_paragraph(name)
        header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = header.runs[0]
        run.font.size = Pt(18)
        run.bold = True

        contact = doc.add_paragraph(f"{email} | {phone} | {location}")
        contact.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        doc.add_paragraph()

        # Summary
        if summary:
            doc.add_heading("Professional Summary", level=1)
            doc.add_paragraph(summary)

        # Skills (bullet points)
        if skills:
            doc.add_heading("Key Skills", level=1)
            for skill in skills.split(","):
                doc.add_paragraph(skill.strip(), style="List Bullet")

        # Experience (bullet points for descriptions)
        if experience:
            doc.add_heading("Work Experience", level=1)
            for exp in experience:
                para = doc.add_paragraph()
                run = para.add_run(
                    f"{exp.get('title','')} — {exp.get('company','')} ({exp.get('duration','')})"
                )
                run.bold = True
                if exp.get("description", ""):
                    for line in exp["description"].split("\n"):
                        if line.strip():
                            doc.add_paragraph(line.strip(), style="List Bullet")

        # Education
        if education:
            doc.add_heading("Education", level=1)
            for edu in education:
                para = doc.add_paragraph()
                run = para.add_run(
                    f"{edu.get('degree','')}, {edu.get('institution','')} ({edu.get('year','')})"
                )
                run.bold = True

        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f"attachment; filename={name.replace(' ','_')}_resume.docx"
            },
        )
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)


# Run locally
if __name__ == "__main__":
    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True)
